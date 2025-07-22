# app/services/extractor.py
#
# Required libraries:
# pip install PyMuPDF python-docx beautifulsoup4

import os
import fitz  # PyMuPDF
from docx import Document
import email
from email.message import Message
from bs4 import BeautifulSoup
from pypdf import PdfReader

# --- Extractor Functions ---


def extract_from_pdf(filepath: str) -> str:
    try:
        output = []
        reader = PdfReader(f"{filepath}")
        for page in reader.pages:
            text = page.extract_text()
            output.append(text)
        return "".join(output)
    except Exception as e:
        print(f"Error processing PDF {filepath}: {e}")
        return ""


def extract_text_from_pdf(filepath: str) -> str:
    """Extracts text from a PDF file using the fast PyMuPDF library."""
    try:
        with fitz.open(filepath) as doc:
            return "".join(page.get_text() for page in doc)
    except Exception as e:
        print(f"Error processing PDF {filepath}: {e}")
        return ""


def extract_text_from_docx(filepath: str) -> str:
    """Extracts text from paragraphs AND tables in a DOCX document."""
    try:
        doc = Document(filepath)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        return "\n\n".join(full_text)  # Use double newline to separate elements
    except Exception as e:
        print(f"Error processing DOCX {filepath}: {e}")
        return ""


def extract_text_from_eml(filepath: str) -> str:
    """Extracts body from email, preferring plain text but falling back to HTML."""
    try:
        with open(filepath, "rb") as f:
            msg = email.message_from_binary_file(f)

        plain_text, html_text = "", ""
        for part in msg.walk():
            if part.get_content_type() == "text/plain" and not plain_text:
                try:
                    plain_text = part.get_payload(decode=True).decode(
                        part.get_content_charset() or "utf-8"
                    )
                except:
                    continue
            elif part.get_content_type() == "text/html" and not html_text:
                try:
                    html_text = part.get_payload(decode=True).decode(
                        part.get_content_charset() or "utf-8"
                    )
                except:
                    continue

        if plain_text.strip():
            return plain_text
        if html_text.strip():
            return BeautifulSoup(html_text, "html.parser").get_text(
                separator="\n", strip=True
            )
        return ""
    except Exception as e:
        print(f"Error processing EML {filepath}: {e}")
        return ""


# --- Main Dispatcher ---

# Using a dictionary is cleaner and more extensible than if/elif chains
EXTRACTORS = {
    ".pdf": extract_from_pdf,
    ".docx": extract_text_from_docx,
    ".eml": extract_text_from_eml,
}


def extract_text(filepath: str) -> str:
    """
    Extracts text from a supported file by dispatching to the correct function.
    """
    ext = os.path.splitext(filepath)[1].lower()
    extractor_func = EXTRACTORS.get(ext)

    if not extractor_func:
        raise ValueError(f"Unsupported file extension: {ext}")

    return extractor_func(filepath)
